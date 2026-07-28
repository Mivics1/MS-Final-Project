#!/usr/bin/env python3
"""
Build the FINAL CO4011 dissertation (all chapters, real results woven in).

This supersedes the earlier results-independent draft. Every number, table and
figure below is produced by the artefact pipeline on the real corpus
(data/emails.csv: balanced real-world Nazario + known AI-generated phishing). Run the
pipeline and the two analysis scripts first, then run this to assemble the .docx:

    .venv/bin/python -m src.pipeline --data data/emails.csv --out results/
    PYTHONPATH=. .venv/bin/python scripts/supplementary_analysis.py
    PYTHONPATH=. .venv/bin/python scripts/make_figures.py
    .venv/bin/python scripts/build_final_dissertation.py

Author: Agboola Michael Daramola
Module: CO4011 Master's Project
"""
from __future__ import annotations

import json
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ART = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RES = os.path.join(ART, "results")
OUT = os.path.join(os.path.dirname(ART), "MS-Final-Dissertation", "Dissertation_Final.docx")

# --- load the real numbers so the prose and tables cannot drift from them ---
results = {r["model"]: r for r in json.load(open(os.path.join(RES, "results.json")))}
supp = json.load(open(os.path.join(RES, "supplementary.json")))
grouped = json.load(open(os.path.join(RES, "grouped_split.json")))
corpus = json.load(open(os.path.join(ART, "data", "corpus_summary.json")))


def f(model, key):
    return results[model][key]


# ---- derived numbers used in the results prose (read from the artefacts) ----
cm_lr = results["logreg"]["confusion_matrix"]
cm_rf = results["random_forest"]["confusion_matrix"]
naz_kept = sum(v["kept_after_clean"] for v in corpus["nazario"].values())
naz_msgs = sum(v["messages"] for v in corpus["nazario"].values())
naz_excl = sum(v.get("excluded_not_credential_harvesting", 0) for v in corpus["nazario"].values())
ai_excl = corpus["eze_shamir"].get("excluded_not_credential_harvesting", 0)
ai_kept = corpus["eze_shamir"]["kept_after_clean"]
per_class = corpus["per_class_after_balancing"]
n_test = sum(cm_lr[0]) + sum(cm_lr[1])
# partition sizes taken from the pipeline's own splitter, so the reported
# figures cannot drift from the split the experiment actually used
from src import data_loader as _dl  # noqa: E402
_split = _dl.stratified_split(_dl.load_csv(os.path.join(ART, "data", "emails.csv")), seed=42)
n_train, n_val = len(_split.y_train), len(_split.y_val)
assert len(_split.y_test) == n_test, "test size disagrees with results.json"
audit = json.load(open(os.path.join(RES, "screen_audit.json")))

# Cite the actual repository state this document was generated from, rather than
# a hand-copied string that can drift.
RELEASE_TAG = "v1.2-reviewed"
try:
    import subprocess
    COMMIT_HASH = subprocess.run(
        ["git", "rev-parse", "--short", "HEAD"], cwd=ART,
        capture_output=True, text=True, check=True).stdout.strip()
except Exception:
    COMMIT_HASH = "(see repository release page)"
cm = supp["class_means"]
abl = supp["ablation"]
agree = supp["importance_agreement"]
gs = grouped["grouped_split"]
gsc = grouped["cluster_stats"]
s2022 = grouped["sensitivity_2022_only"]


doc = Document()

# ============================ base styling ============================
for s in doc.sections:
    s.page_height = Inches(11.69)
    s.page_width = Inches(8.27)  # A4
    s.left_margin = s.right_margin = Inches(1.0)
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)
for hs, sz in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)]:
    st = doc.styles[hs]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)


def para(t, justify=True, after=6):
    p = doc.add_paragraph(t)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(after)
    return p


def h1(t):
    return doc.add_heading(t, 1)


def h2(t):
    return doc.add_heading(t, 2)


def bullet(t):
    doc.add_paragraph(t, style="List Bullet")


def numbered(t):
    doc.add_paragraph(t, style="List Number")


def page_break():
    doc.add_page_break()


def lead(label, body):
    """A bold lead-in run followed by justified body text."""
    p = doc.add_paragraph()
    r = p.add_run(label + " ")
    r.bold = True
    p.add_run(body)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def caption(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.italic = True
    r.font.size = Pt(9.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    stripped = t.lstrip()
    if stripped.startswith("Figure"):
        FIGLIST.append(t)
    elif stripped.startswith("Table"):
        TBLLIST.append(t)


def figure(path, width_in, cap):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    caption(cap)


def add_hyperlink(paragraph, url, text):
    """Insert a real, clickable external hyperlink into a paragraph."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rpr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run)
    paragraph._p.append(link)
    return link


# --- page numbering / fields --------------------------------------------------
FIGLIST: list[str] = []
TBLLIST: list[str] = []


def _add_page_field(paragraph):
    run = paragraph.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    i = OxmlElement("w:instrText"); i.set(qn("xml:space"), "preserve"); i.text = " PAGE "
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(i); run._r.append(e)


def set_page_numbering(section, fmt, start=None):
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType"); sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))


def footer_page_number(section, blank_first=False):
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_field(p)
    if blank_first:
        section.different_first_page_header_footer = True


def insert_after(anchor_par, text, size=10, bold=False):
    new_p = OxmlElement("w:p")
    anchor_par._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    par = Paragraph(new_p, anchor_par._parent)
    r = par.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return par


def _set_cell(cell, text, bold=False, size=9.5, fill=None, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                   "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "right": WD_ALIGN_PARAGRAPH.RIGHT}[align]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if fill:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)


def _cant_split_row(row):
    """Stop a table row from breaking across a page."""
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    trPr.append(cs)


def _repeat_header(row):
    """Repeat this row as a header on each page the table spans."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def table(headers, rows, widths=None, aligns=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aligns = aligns or ["left"] + ["center"] * (len(headers) - 1)
    for j, htext in enumerate(headers):
        _set_cell(t.rows[0].cells[j], htext, bold=True, fill="1F2A37", align=aligns[j])
        t.rows[0].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            _set_cell(cells[j], str(val), align=aligns[j])
    if widths:
        for j, wd in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(wd)
    # keep every row intact, repeat the header, and bind rows together so the
    # whole table (and the caption that follows) stays on one page where it fits
    for r in t.rows:
        _cant_split_row(r)
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.keep_with_next = True
    _repeat_header(t.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================ TITLE PAGE ============================
for _ in range(3):
    doc.add_paragraph()
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Stylometric Detection of AI-Generated Phishing")
r.bold = True
r.font.size = Pt(22)
sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run("Distinguishing Known AI-Generated Credential-Harvesting Emails "
               "from Real-World Phishing")
r.italic = True
r.font.size = Pt(13)
for _ in range(3):
    doc.add_paragraph()
for label, val in [("Author", "Agboola Michael Daramola"),
                   ("Student ID", "[Student ID]"),
                   ("Supervisor", "Dr Radu Negoescu"),
                   ("Module", "CO4011 Master's Project (Level 7)"),
                   ("Programme", "MSc Cybersecurity with AI"),
                   ("Institution", "University of Lancashire")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p.add_run(f"{label}: ")
    rb.bold = True
    p.add_run(val)
for _ in range(3):
    doc.add_paragraph()
d = doc.add_paragraph()
d.alignment = WD_ALIGN_PARAGRAPH.CENTER
d.add_run("Deliverable 2 — Project Report").italic = True
page_break()

# ============================ ABSTRACT ============================
h1("Abstract")
para(
    "Large language models have removed the historical trade-off between the scale "
    "and the quality of phishing, allowing attackers to mass-produce fluent, "
    "personalised credential-harvesting emails that evade the spelling and grammar "
    "cues on which many defences once relied. Building on Eze and Shamir (2024), who "
    "showed that stylistic text analysis can separate AI-generated phishing from human "
    "email classes, this project is a focused replication and extension that asks a "
    "single question within one narrow attack type: to what extent do interpretable "
    "stylometric features separate known AI-generated phishing from real-world "
    "phishing, and which markers carry the signal? A balanced corpus of "
    f"{corpus['total']} credential-harvesting emails was assembled from two published "
    "research sources — "
    f"{corpus['per_class_after_balancing']} real-world phishing messages from the "
    "Nazario corpus (2022–2024) and "
    f"{corpus['per_class_after_balancing']} known AI-generated messages from the Eze "
    "and Shamir (2024) dataset — after a documented, rule-based credential-harvesting "
    "screen. Each email body was reduced to a vector of 95 interpretable stylometric "
    "and lexical features, computed with no access to the generating model. "
    "Logistic-regression, Random Forest and XGBoost classifiers were trained through a "
    "fully seeded, leakage-controlled pipeline and evaluated on an untouched held-out "
    "test set. All three models exceeded the target of F1 = 0.90 pre-specified in the "
    "approved proposal: the logistic-regression model achieved an F1 of "
    f"{f('logreg','f1'):.3f} (95% CI [{f('logreg','f1_ci_low'):.3f}, "
    f"{f('logreg','f1_ci_high'):.3f}]) and a ROC-AUC of {f('logreg','roc_auc'):.3f}, "
    "and the tree ensembles reached F1 = "
    f"{f('random_forest','f1'):.3f}. The signal is dominated by readability and "
    "lexical regularity: AI-generated emails are markedly harder to read (mean Flesch "
    f"Reading Ease {supp['class_means']['flesch_reading_ease']['llm_mean']:.1f} versus "
    f"{supp['class_means']['flesch_reading_ease']['human_mean']:.1f}), use longer "
    "words, and deploy exclamation marks and urgency vocabulary more consistently. "
    "Removing the two most obvious structural cues changed performance little "
    f"(F1 {abl['logreg']['f1']:.3f}), and performance did not fall under a "
    "campaign-level (grouped) split that kept every detected template cluster within a "
    "single partition, indicating the result is not driven solely by those cues or by "
    "template leakage; broader source-level confounding nonetheless remains possible. "
    "A manual audit of the corpus screening rule is reported, including a "
    "source-asymmetry defect it uncovered and corrected. The study concludes that "
    "stylometry is a viable and interpretable detection signal for AI-generated "
    "phishing, subject to the limitations that the AI class comes from a single "
    "published generation dataset, that authorship of the real-world class is not "
    "verified, and that the detector may be vulnerable to deliberate paraphrasing.",
    after=6)
page_break()

# ============================ TOC ============================
h1("Table of Contents")
p = doc.add_paragraph()
run = p.add_run()
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
cached = OxmlElement("w:r")
ct = OxmlElement("w:t")
ct.text = "The table of contents updates automatically on opening (Word will prompt); " \
          "otherwise select all and press F9."
cached.append(ct)
fld.append(cached)
p._p.append(fld)
page_break()

# ---- List of Figures / List of Tables (filled at the end from collected captions) ----
h1("List of Figures")
LOF_ANCHOR = doc.add_paragraph()
page_break()
h1("List of Tables")
LOT_ANCHOR = doc.add_paragraph()

# ---- section break: prelims (Roman numerals) end here; body (Arabic) begins ----
from docx.enum.section import WD_SECTION
BODY_SECTION = doc.add_section(WD_SECTION.NEW_PAGE)

# ============================ 1. INTRODUCTION ============================
h1("1. Introduction")

para(
    "Phishing remains the single most common route by which attackers gain an initial "
    "foothold in an organisation, and the harvesting of user credentials is its "
    "dominant objective. For most of its history the economics of phishing imposed a "
    "natural ceiling on quality: a convincing, well-targeted message required human "
    "effort, so an attacker had to trade scale against persuasiveness. The maturation "
    "of large language models (LLMs) has dissolved that trade-off. A capable model can "
    "now be instructed to produce fluent, grammatically clean and context-aware "
    "credential-harvesting emails in seconds, and to emit a fresh variant for every "
    "recipient. Heiding et al. (2024) found that fully automated GPT-4 phishing "
    "outperformed generic control messages (roughly 30–44% click-through) but did not, "
    "on its own, match emails designed with the human expert V-Triad framework (roughly "
    "69–79%); a hybrid of GPT-4 and the V-Triad performed comparably to, and in some "
    "conditions better than, the human-only method — all at a fraction of the human "
    "effort. The threat is therefore one of scale and cost-efficiency as much as raw "
    "persuasiveness. The problem this project addresses is the detection of "
    "machine-authored phishing, and specifically whether the writing style of such "
    "messages betrays their non-human origin.")

para(
    "The problem decomposes into three connected parts. First, the industrialisation "
    "of phishing by LLMs raises both the volume and the average quality of attacks, "
    "eroding the grammatical and spelling errors on which users and some filters have "
    "traditionally relied. Second, because a generative model can produce a unique "
    "message for each victim, defences that match known-bad signatures or blocklisted "
    "content lose much of their purchase (Eze and Shamir, 2024). Third, and least "
    "understood, is the question of authorship: it is not yet established whether the "
    "stylistic regularities of machine-generated prose are strong and stable enough to "
    "act as a reliable detection signal for phishing specifically, rather than for "
    "long-form text in general. This project isolates and investigates that third "
    "component.")

para(
    "Several standards and legal instruments frame the work. Email corpora may contain "
    "personal data — names, addresses and other identifiers — so the project treats its "
    "source data as personal data under the UK General Data Protection Regulation and "
    "the Data Protection Act 2018 unless effective anonymisation can be demonstrated, "
    "and the principles of data minimisation and storage limitation govern how it is "
    "handled. The Computer Misuse Act 1990 makes the "
    "boundary between defensive research and offensive capability legally significant, "
    "and this project sits firmly on the defensive side: it detects, and never "
    "deploys, phishing. More broadly the work aligns with the risk-management framing "
    "of ISO/IEC 27001 and with the anti-phishing guidance issued by national bodies "
    "such as the UK National Cyber Security Centre. These considerations are developed "
    "in the Methodology.")

para(
    "The significance of the problem follows from its trajectory. Industry telemetry "
    "such as the Anti-Phishing Working Group's trends reports and the annual Verizon "
    "Data Breach Investigations Report consistently place phishing and stolen "
    "credentials among the leading causes of security breaches (APWG, 2025; Verizon, "
    "2025). If the problem is solved, even partially, a stylometric detector could add "
    "a lightweight, content-based layer to existing mail-filtering and "
    "security-operations tooling, flagging messages whose style is characteristic of "
    "machine generation. If it is not, defenders face an adversary that can scale "
    "high-quality, personalised social engineering almost without limit while the cues "
    "that once exposed phishing steadily vanish. The asymmetry of these outcomes "
    "motivates a focused, rigorous study.")

para(
    "This project works towards a solution through a controlled, reproducible "
    "machine-learning experiment. It deliberately restricts its scope to one attack "
    "type (credential-harvesting emails), one family of features (stylometric and "
    "lexical), and one binary decision (known AI-generated versus real-world phishing), "
    "so that the outcome is a definitive answer to a single question rather than a "
    "broad but shallow survey.")

lead("Aim.",
     "To determine the extent to which stylometric and lexical features enable "
     "machine-learning classifiers to distinguish known AI-generated "
     "credential-harvesting phishing emails from real-world (Nazario) phishing "
     "emails, and to identify which "
     "features are most discriminative.")

p = doc.add_paragraph()
p.add_run("Objectives.").bold = True
numbered("Critically review the literature on AI-generated phishing, "
         "machine-generated-text detection and stylometric authorship attribution, "
         "establishing the boundary between what is known and what is not.")
numbered("Assemble a balanced, fully documented binary corpus of real-world and known "
         "AI-generated credential-harvesting emails from published research sources.")
numbered("Specify and implement a stylometric and lexical feature set that is "
         "interpretable as well as discriminative.")
numbered("Train and tune Random Forest and XGBoost classifiers, with a standardised "
         "logistic-regression baseline, using a reproducible, seed-fixed pipeline.")
numbered("Evaluate all models on a held-out test set using precision, recall and F1 "
         "with bootstrap confidence intervals, against a target F1 of at least 0.90.")
numbered("Rank and interpret feature importance to identify the stylistic markers "
         "that most reveal machine authorship.")
numbered("Test whether performance depends on selected corpus-specific structural "
         "cues, and assess robustness to campaign-level template leakage.")
page_break()

# ============================ 2. STATE OF THE ART ============================
h1("2. State of the Art")

h2("2.1 The threat: effectiveness of LLM-generated phishing")
para(
    "The starting point for this review is evidence that machine-generated phishing is "
    "a real and potent threat rather than a speculative one. Heiding et al. (2024) "
    "compared phishing produced automatically by GPT-4 with messages designed using the "
    "human expert V-Triad framework. Fully automated GPT-4 emails beat generic controls "
    "but did not, alone, reach the click-through rates of the human-designed messages; "
    "it was the hybrid of GPT-4 and the V-Triad that matched or exceeded the human-only "
    "method. Their results matter here in two ways: they establish the motivation — "
    "cheap, scalable phishing that is already competitive once lightly steered by a "
    "human — and they show that fluency and correctness, the properties LLMs excel at, "
    "are central to the threat. This suggests, in turn, that any residual stylistic "
    "signature of machine generation is where a defender should look.")

h2("2.2 Detecting machine-generated text in general")
para(
    "A large and fast-moving body of work addresses the detection of AI-generated text "
    "irrespective of domain. Three families of method have emerged: supervised "
    "classifiers that learn to separate human from machine text from labelled "
    "examples; zero-shot statistical methods that exploit intrinsic properties of "
    "model output; and watermarking, which embeds a detectable signal at generation "
    "time. The most influential zero-shot method, DetectGPT (Mitchell et al., 2023), "
    "observes that machine-generated passages tend to occupy regions of negative "
    "curvature in a model's log-probability surface and detects them by measuring how "
    "sharply likelihood changes under small perturbations; it improved detection of "
    "long-form machine text markedly over earlier perplexity baselines. However, this "
    "family has a well-documented fragility. Krishna et al. (2023) showed that "
    "paraphrasing machine output with a dedicated model collapses the accuracy of "
    "DetectGPT, watermarking and commercial detectors alike — DetectGPT's detection "
    "rate fell from roughly 70% to under 5% at a fixed false-positive rate — while "
    "meaning is preserved. Two implications follow for the present work: detectors "
    "that need the generating model's probabilities are impractical for an email "
    "defender who does not control the attacker's model, and robustness must be "
    "treated as an explicit limitation rather than assumed away.")
para(
    "The supervised family is the closest cousin of this project's approach and repays "
    "closer attention. Early large-model release studies established both its promise "
    "and its ceiling. Solaiman et al. (2019), reporting on the staged release of GPT-2, "
    "found that a fine-tuned classifier could detect that model's output with high "
    "accuracy but that performance degraded as the generator grew more capable and as "
    "sampling strategies changed — an early warning that detectability is a moving "
    "target tied to a specific generator. Zellers et al. (2019), with Grover, made the "
    "sharper argument that the strongest detector of neural text is often a model from "
    "the same family as the generator, which is precisely the assumption an email "
    "defender cannot make in practice. Interactive tools such as GLTR (Gehrmann et al., "
    "2019) exposed the per-token statistical fingerprints of machine text to human "
    "reviewers, reinforcing that the signal is real but that its features are "
    "generator-dependent. Taken together, this strand motivates the design choice at "
    "the heart of the present work: rather than depend on the generator's own "
    "probabilities or a same-family reference model, it uses generator-independent "
    "stylometric features that a defender can compute from message text alone, "
    "trading some theoretical power for practicality and interpretability.")
para(
    "The watermarking strand deserves separate mention because it addresses a different "
    "point in the pipeline. Rather than detect text after the fact, Kirchenbauer et al. "
    "(2023) proposed biasing a model's token sampling towards a secret 'green list' at "
    "generation time, so that generated text carries a statistically detectable "
    "signature invisible to a human reader. This is elegant where the defender "
    "controls, or can compel, the generator — but it is inapplicable to the phishing "
    "setting, where the attacker chooses the model and will simply use an unwatermarked "
    "one. Worse, the same paraphrasing attack that breaks post-hoc detectors also "
    "weakens watermarks (Krishna et al., 2023). The overall picture from the "
    "machine-text-detection literature is therefore an arms race in which every "
    "detection signal has a corresponding evasion, and in which the defender's realistic "
    "assumptions — no model access, no cooperation from the generator — rule out the most "
    "powerful methods. That constraint, rather than a preference, is what makes a "
    "transparent stylometric approach worth investigating for the specific, bounded "
    "problem of phishing email.")

h2("2.3 Stylometry and authorship attribution")
para(
    "The idea that writing style carries a measurable, author-specific signature "
    "predates modern machine learning by decades. The seminal study of Mosteller and "
    "Wallace (1964) attributed the disputed Federalist Papers by analysing the "
    "frequencies of common function words — articles, prepositions, conjunctions — on "
    "the grounds that such words are used unconsciously and are largely independent of "
    "topic. Function-word frequencies have since become a standard feature in "
    "authorship attribution precisely because of this topic-independence and "
    "resistance to deliberate manipulation (Stamatatos, 2009). The relevance here is "
    "direct: if human and machine authors differ in their unconscious stylistic habits "
    "— sentence-length regularity, vocabulary richness, punctuation and function-word "
    "distributions — then a stylometric feature set offers an interpretable, "
    "generator-independent alternative to probability-based detectors. This project transfers "
    "the authorship-attribution toolkit from its classical human-versus-human setting "
    "to the human-versus-machine problem within the narrow domain of phishing email.")

h2("2.4 Detecting AI-generated phishing specifically")
para(
    "A smaller but growing literature targets AI-generated phishing directly, and Eze "
    "and Shamir (2024) are the direct precedent for this work — they supply its AI "
    "class and its method. They analysed a corpus of AI-generated phishing with UDAT, a "
    "stylistic approach built on characteristics such as punctuation, word length, "
    "readability, parts of speech and repetition, ranked the most informative "
    "descriptors (reporting differences in word length, pronoun and verb frequency, "
    "lexical diversity and sentence length), and reported two-way classification between "
    "AI-generated and human scam email with high accuracy. Koide et al. (2024), with "
    "ChatSpamDetector, took a complementary route, prompting an LLM to classify emails "
    "and explain its decisions. Together these studies establish that stylistic "
    "detection of AI-generated phishing is tractable and that several stylistic "
    "descriptors are informative. This project is therefore not the first demonstration "
    "of style-based detection; it is a focused replication and extension.")
para(
    "The extension is specific. Eze and Shamir evaluate over a mix of email types and "
    "against legitimate and human scam email, so a classifier can lean on topic or "
    "structural signals as well as style, and the reported accuracy is treated as the "
    "endpoint. The present work holds several variables fixed instead: it compares a "
    "single attack type (credential harvesting) against real-world phishing rather than "
    "legitimate mail, restricts the evidence to body-text style through a transparent "
    "95-feature implementation, adds seed-fixed evaluation with bootstrap confidence "
    "intervals, ranks the responsible features by two methods, and adds an explicit "
    "ablation of structural cues plus a campaign-level grouped split to probe whether "
    "the accuracy is an artefact of corpus construction. The contribution is thus a "
    "narrower, better-characterised and more reproducible account of a signal earlier "
    "work already showed exists.")

h2("2.5 Classical phishing detection and its blind spot")
para(
    "For completeness the review situates the work against conventional phishing "
    "detection, which has historically relied on features drawn from message headers, "
    "embedded URLs and sender reputation, together with content-based natural-language "
    "features. The canonical work in this tradition, Fette, Sadeh and Tomasic (2007), "
    "trained a classifier ('PILFER') on ten hand-engineered structural features — the "
    "number of links, the presence of IP-based URLs, mismatched href domains, and "
    "similar signals — and reported detection well above 90%, and Abu-Nimeh et al. "
    "(2007) compared six machine-learning methods on a comparable feature set, "
    "establishing the experimental template the field still follows. These methods "
    "remain valuable, but two things date them for the present purpose. They largely "
    "predate the generative-AI threat, and their most powerful features are structural "
    "artefacts of the message envelope rather than properties of its prose — the very "
    "features an LLM-generated message may lack or randomise. More fundamentally they "
    "answer the question 'is this message malicious?' rather than 'was this message "
    "written by a machine?'. The authorship question is orthogonal to, and "
    "complementary with, the maliciousness question, and it is this gap the project "
    "occupies. It also deliberately avoids leaning on envelope artefacts: Section 6.4 "
    "shows the result holds when the two most envelope-like features are removed.")

h2("2.6 Research methodologies in the field and their implications")
para(
    "Methodologically, the reviewed work is dominated by quantitative, experimental "
    "studies: authors construct or obtain a labelled corpus, engineer or learn "
    "features, train and tune classifiers, and report standard classification metrics "
    "on held-out data. Recurring weaknesses are visible across the field — modest or "
    "single-source datasets, reliance on a single generator model, thin reporting of "
    "statistical uncertainty, and incomplete reproducibility. These observations "
    "directly shaped the methodology adopted here: a quantitative experimental design "
    "is appropriate and well-precedented, but it must be strengthened with explicit "
    "reproducibility guarantees, confidence intervals on the headline metric, and an "
    "honest, up-front statement of the single-generator limitation.")

h2("2.7 The gap this project addresses")
para(
    "In summary, it is known that LLM-generated phishing is effective when combined with "
    "human design (Heiding et al., 2024), that machine text is in principle detectable "
    "but that probability-based detectors are fragile under paraphrasing and require "
    "model access (Mitchell et al., 2023; Krishna et al., 2023), that stylometry can "
    "attribute authorship from unconscious stylistic habits (Mosteller and Wallace, "
    "1964; Stamatatos, 2009), and — most directly — that stylistic analysis can already "
    "separate AI-generated phishing from human email and identify informative "
    "descriptors (Eze and Shamir, 2024; Koide et al., 2024). What remains less well "
    "characterised is how that signal behaves under tighter controls: within a single "
    "attack type, against real-world phishing rather than legitimate mail, with a "
    "transparent feature set whose contributions are ranked, and with explicit tests of "
    "whether the performance survives structural-cue ablation and campaign-level "
    "leakage control. Providing that narrower, reproducible characterisation — an "
    "extension of Eze and Shamir rather than a first demonstration — is the contribution "
    "of this project.")
page_break()

# ============================ 3. METHODOLOGY ============================
h1("3. Methodology")

h2("3.1 Research design")
para(
    "The project adopts a quantitative, experimental research design. The research "
    "question is fundamentally measurable — it asks to what extent a classifier can "
    "separate two classes and which features drive that separation — so a quantitative "
    "approach is more appropriate than a qualitative one, and it is consistent with "
    "the dominant methodology in the field (Section 2.6). The design is comparative: "
    "several models are trained on a common feature representation and evaluated under "
    "identical conditions, so that differences in performance can be attributed to the "
    "models and features rather than to inconsistent experimental treatment.")

h2("3.2 Data collection and sampling")
para(
    "The study uses secondary data only, which keeps it reproducible and avoids the "
    "ethical hazard of generating fresh attack content. The two classes are named "
    "precisely to reflect what the sources actually guarantee. The first is real-world "
    "phishing from the Nazario corpus, an established public collection; its 2022–2024 "
    "archives were chosen so the class is contemporaneous with the AI class, avoiding an "
    "era-of-writing confound. This class is operationally treated as human-origin, but "
    "the corpus is hand-classified as phishing rather than verified for authorship, so "
    "it is labelled 'real-world' rather than 'human-written' (see Section 3.7). The "
    "second is known AI-generated phishing from the published Eze and Shamir (2024) "
    "dataset, used rather than produced afresh. Both sources are reduced to email body "
    "text so the classifier sees the same kind of evidence for each class, and a "
    "documented, rule-based credential-harvesting screen is applied to both so the "
    "corpus matches the narrow scope: a message is retained only if it both references "
    "an account, credential or verification subject and solicits an action through a "
    "link, form or reply. The number excluded by this screen is reported per source in "
    "Section 5.1 and Appendix B. The two classes are then balanced by construction to "
    "prevent the classifier from exploiting prior class frequency, and the corpus is "
    "de-duplicated before any splitting. The data are partitioned by stratified sampling "
    "under a fixed random seed. A stratified 70/15/15 training/validation/test split is "
    "produced; hyperparameter selection uses five-fold cross-validation inside the "
    "training partition only, so the 15% validation partition is held in reserve and "
    "not used for tuning or reporting, and all held-out results are computed on the "
    "untouched 15% test partition. Exact partition counts are given in Section 5.3.")

h2("3.3 Analysis methods")
para(
    "Each email body is converted into a fixed-length vector of interpretable "
    "stylometric and lexical features (specified in Chapter 4): measures of length, "
    "vocabulary richness, readability, punctuation and orthography, structural cues, "
    "and function-word frequencies. Logistic regression, Random Forest and XGBoost "
    "classifiers are trained on these vectors, the first as a standardised baseline "
    "floor and the latter two as the primary models. Hyperparameters are selected by "
    "five-fold cross-validated grid search on the training data only. Performance is "
    "analysed with precision, recall and F1 for the positive (AI-generated) class, "
    "supplemented by ROC-AUC and confusion matrices, and the uncertainty of the "
    "headline F1 is quantified with a percentile bootstrap confidence interval. "
    "Feature importance is analysed with both impurity-based and permutation-based "
    "methods so the interpretation does not rest on a single, potentially biased "
    "estimator, and per-class feature means are inspected to establish the direction "
    "of each signal.")
para(
    "The choice of metric deserves justification because it is where several reviewed "
    "studies were weakest. Accuracy alone is a poor summary even for a balanced corpus "
    "because it hides the trade-off between the two error types that matter to a "
    "defender — a missed machine-phishing email versus a wrongly flagged human one — so "
    "precision, recall and their harmonic mean, F1, are reported for the positive "
    "class, with ROC-AUC to summarise performance across all decision thresholds. "
    "Because a single point estimate on one test split can mislead, the headline F1 is "
    "accompanied by a percentile bootstrap 95% confidence interval computed by "
    "resampling the test predictions, which quantifies how much the result would be "
    "expected to vary under resampling and guards against over-claiming from a lucky "
    "split. A success threshold of F1 = 0.90, pre-specified in the approved proposal, "
    "fixes the bar before the results are seen so that the evaluation cannot be "
    "retrofitted to whatever the models happened to achieve. (This is a "
    "proposal-stated target, not a formal public preregistration.)")

h2("3.4 Objective completion criteria")
para(
    "Each objective has an explicit test of completion. Objective 1 is complete when "
    "the literature review establishes the known/unknown boundary with adequate "
    "coverage; Objective 2 when a balanced, de-duplicated, documented corpus exists; "
    "Objective 3 when the feature extractor produces a stable vector and passes its "
    "unit tests; Objective 4 when the models are trained with tuned hyperparameters "
    "and a fixed seed; Objective 5 when all models are evaluated on the untouched test "
    "set with confidence intervals; Objective 6 when a ranked, interpreted list of "
    "discriminative features has been produced; and Objective 7 when structural-feature "
    "ablation and grouped splitting have quantified the extent to which selected corpus "
    "cues and detected template clusters influence performance. This progression is "
    "staged so that progress is observable and reportable at each step.")

h2("3.5 Evaluation process")
para(
    "Evaluation proceeds in two layers. The first is quantitative and answers whether "
    "the aim has been met: does the best model reach the target F1 of 0.90 on held-out "
    "data, and how does it compare with the simple baselines and with published "
    "results? The second is interpretive and answers the second half of the aim: which "
    "stylistic markers most reveal machine authorship, and are they consistent with "
    "the theoretical expectation that machine prose is more fluent and regular? "
    "Critically, the evaluation also reflects on its own limits — in particular, what "
    "the chosen metrics cannot reveal about generalisation to generator models not "
    "seen during training.")

h2("3.6 Ethical considerations")
para(
    "The work is defensive dual-use research and is treated accordingly. No message is "
    "ever sent to any recipient, no real individual or organisation is targeted, and no "
    "new phishing content is generated for release. Email corpora may contain personal "
    "data — names, addresses and other identifiers — so under UK GDPR and the Data "
    "Protection Act 2018 the project treats the source data as personal data unless "
    "effective anonymisation can be demonstrated. Accordingly it relies on pre-existing "
    "public research corpora, collects no additional personal data, stores data on an "
    "encrypted device, excludes raw phishing text from the public artefact repository "
    "(only aggregate statistics and derived features are published), and deletes the "
    "working data at project end in line with university policy. The Nazario corpus is "
    "used under its CC BY 4.0 licence with attribution to Jose Nazario. No human "
    "participants are involved, so no participant information sheet or consent form is "
    "required; on that basis the module's supervisor-led ethics process confirmed the "
    "project before any data handling began, and no separate formal approval number was "
    "issued for this secondary-data study.")

h2("3.7 Limitations")
para(
    "Four limitations are acknowledged in advance. First, the AI class originates from a "
    "single published AI-generation dataset produced through the DeepAI text-generation "
    "service, so the findings speak to detecting that dataset's output; the exact "
    "underlying model and version are not established by the source, and generalisation "
    "to other generators is untested. Second, the real-world (Nazario) class is hand-"
    "classified as phishing but not verified as human-authored, and because its archives "
    "overlap the period of public LLM availability, some AI-assisted contamination "
    "cannot be ruled out; a 2022-only sensitivity analysis (Section 6.5) partially "
    "addresses this. Third, the study is limited to English-language credential-"
    "harvesting email. Fourth, consistent with Krishna et al. (2023), a stylometric "
    "detector may be vulnerable to deliberate paraphrasing or 'humanisation' of machine "
    "output — an avenue for future work rather than a defect of the present study. These "
    "limitations are revisited against the actual results in Chapters 6 and 7.")
page_break()

# ============================ 4. DESIGN ============================
h1("4. Design")

h2("4.1 Requirements analysis")
para(
    "The functional requirements of the artefact follow directly from the objectives. "
    "The system must (F1) ingest a labelled corpus of emails and clean and de-duplicate "
    "it; (F2) extract a fixed, interpretable stylometric and lexical feature vector "
    "from each message; (F3) train and tune the baseline, Random Forest and XGBoost "
    "models; (F4) evaluate every model on a held-out test set and report precision, "
    "recall, F1, ROC-AUC and confidence intervals; and (F5) produce a ranked, "
    "interpreted feature-importance analysis. The non-functional requirements matter "
    "just as much for a research artefact: the pipeline must be reproducible (every "
    "stochastic step seeded), transparent (features must be human-interpretable so "
    "results can be explained and defended in the viva), efficient enough to run on a "
    "laptop, and testable (feature extraction and metrics covered by unit tests).")
para(
    "Interpretability was elevated from a nicety to a hard requirement, and this shaped "
    "the whole design. A more powerful detector could in principle be built by "
    "fine-tuning a transformer such as BERT (Devlin et al., 2019) directly on the raw "
    "text, and the literature confirms such models are strong text classifiers. But a "
    "transformer's decision is opaque, and the second half of the aim — identifying "
    "which stylistic markers reveal machine authorship — cannot be answered by a model "
    "whose features are uninterpretable latent dimensions. Choosing an explicit, named "
    "feature set over a learned representation is therefore not a limitation forced by "
    "compute but a decision that follows directly from the research question, and it is "
    "what lets Chapter 6 report not just that style discriminates but how and in which "
    "direction.")

h2("4.2 System architecture")
para(
    "The artefact is organised as a linear, modular pipeline in which each stage has a "
    "single responsibility and a well-defined interface to the next: a data-loading "
    "and splitting module, a feature-extraction module, a training module, an "
    "evaluation module, and an interpretation module, coordinated by a single "
    "orchestration entry point (Figure 4.1). This separation makes the system easy to "
    "test, to reason about, and to extend — adding a SHAP-based interpretation or an "
    "additional model affects only one module. Deliberately, the feature-extraction "
    "module depends only on the standard library and numpy, so the scientific core "
    "runs in any environment and can be unit-tested without the heavier "
    "machine-learning dependencies.")
figure(os.path.join(RES, "fig_architecture.png"), 6.3,
       "Figure 4.1 — Pipeline architecture. Each stage has a single responsibility "
       "and a fixed interface, and every stochastic step is seeded for reproducibility.")

h2("4.3 Feature-set specification")
para(
    "The feature set is the central design artefact and is specified to be "
    "interpretable as well as discriminative, so the eventual importance ranking is "
    "meaningful. It comprises 95 features in six groups, summarised in Table 4.1: 25 "
    "document-level stylometric and lexical features plus the relative frequencies of "
    "70 English function words. The inclusion of each group is justified by the "
    "literature: function words and vocabulary-richness statistics are the classical "
    "authorship-attribution signal (Mosteller and Wallace, 1964; Stamatatos, 2009), "
    "while regularity of length and readability operationalise the expectation that "
    "machine prose is more uniform and fluent than human writing.")
table(
    ["Feature group", "Count", "Examples", "Rationale"],
    [["Volume & length", "7",
      "char/word/sentence counts; mean & s.d. of word and sentence length",
      "Machine output tends to be more uniform in length"],
     ["Vocabulary richness", "5",
      "type–token ratio, hapax & dis-legomena ratios, Honoré's R, Yule's K",
      "Classical authorship signal; lexical diversity differs by author"],
     ["Readability", "3",
      "Flesch Reading Ease, Flesch–Kincaid grade, syllables/word",
      "Fluency and register of machine prose"],
     ["Punctuation & orthography", "7",
      "comma, period, exclamation, question, uppercase, digit ratios",
      "Stylistic habits and formatting regularity"],
     ["Structural / phishing cues", "3",
      "URL count, email-address count, urgency-word ratio",
      "Interpretable content markers (tested for artefact in §6.4)"],
     ["Function-word frequencies", "70",
      "the, a, of, to, you, is, were, do, our …",
      "Unconscious, topic-independent stylometric fingerprint"]],
    widths=[1.6, 0.55, 2.3, 1.85])
caption("Table 4.1 — The 95-feature stylometric and lexical feature set (see "
        "src/features.py for exact definitions).")

h2("4.4 Requirements prioritisation (MoSCoW)")
para(
    "To manage an intensive timeline the requirements were prioritised. The Must-have "
    "elements are the corpus, the feature extractor, the Random Forest and XGBoost "
    "models, and the core metric report. The Should-have element is a fine-tuned "
    "transformer baseline. The Could-have elements are SHAP-based interpretation and a "
    "cross-generator robustness test. The Won't-have elements, explicitly out of "
    "scope, are any live deployment, multi-language support, and "
    "non-credential-harvesting attack types. This prioritisation guaranteed that even "
    "under time pressure the artefact would still answer the research question, "
    "because the Must-have set alone is sufficient — and, as Chapter 5 records, the "
    "transformer baseline was ultimately cut in line with the proposal's contingency "
    "plan without weakening the answer.")
page_break()

# ============================ 5. IMPLEMENTATION ============================
h1("5. Implementation")
para(
    "This chapter reviews the key problems encountered while building and running the "
    "artefact and how each was overcome, rather than narrating every line of code; the "
    "full source is referenced in Appendix A. Five problems were substantive enough to "
    "shape the result: corpus construction and class comparability, train/test "
    "leakage, feature-extraction edge cases, hyperparameter search under modest "
    "compute, and the decision to cut the transformer baseline. A review of the "
    "testing process closes the chapter.")

h2("5.1 Corpus construction and class comparability")
para(
    "The most consequential engineering was not modelling but data. The two source "
    "corpora arrive in incompatible shapes: the AI data are 865 individual plain-text "
    "files each headed by a synthetic subject line, whereas the real-world data are "
    "large mbox archives of raw MIME messages with headers, HTML parts and "
    "quoted-printable encodings. A naive load would have handed the classifier trivial, "
    "non-stylistic giveaways — MIME boundaries, HTML tags and header keywords appear in "
    "one class and never in the other. The corpus builder therefore parses each mbox "
    "message with the standard-library email parser, prefers the text/plain part and "
    "falls back to HTML stripped to text, strips the synthetic 'Header:' line from the "
    "AI files so both classes are body-only, normalises whitespace and Unicode, and "
    "applies a stopword-ratio heuristic to drop non-English bodies.")
para(
    "A rule-based credential-harvesting screen is then applied to both classes so the "
    "corpus matches the project's stated scope, addressing the risk that the AI source "
    "in particular contains event invitations and other non-credential messages. The "
    "operational definition is fixed in the build script: a message is retained only if "
    "it references an account, credential or verification subject and also solicits an "
    "action through a link, button, form or reply. This is a deterministic "
    "keyword-and-pattern rule, applied identically to both sources and fully "
    "reproducible, but determinism is not correctness — a fixed rule can still make "
    "systematic errors, so it was validated by manual audit (Section 5.2). The screen "
    f"excluded {ai_excl} of the 865 AI messages (retaining {ai_kept}) and {naz_excl} "
    f"real-world messages, from {naz_msgs} raw Nazario messages of which {naz_kept} "
    "passed the screen after cleaning. Exact-duplicate removal on the lower-cased "
    "alphanumeric text then removed reformatted copies of repeated templates, and a "
    f"seeded down-sample balanced the classes to {per_class} emails each "
    f"({corpus['total']} in total). A second comparability decision — restricting the "
    "real-world class to the 2022–2024 archives so it is contemporaneous with the 2024 "
    "AI class — prevents the model from separating the classes on period-specific "
    "vocabulary rather than style. Every one of these decisions lives in the runnable "
    "build script with its counts, which is what makes the corpus reproducible.")

h2("5.2 Validating the credential-harvesting screen")
para(
    "Because the screen determines which messages enter the corpus, its validity "
    "underpins the claim that the study is about credential harvesting specifically, so "
    "it was audited rather than assumed correct. A seeded stratified random sample of "
    "100 messages — 25 retained and 25 excluded from each source — was drawn and each "
    "message manually classified against the operational definition. The audit was "
    "conducted by the author alone; no second rater was available, so no inter-rater "
    "agreement statistic is reported, and the judgements are recorded in the artefact "
    "so that they can be re-examined.")
para(
    "The first round exposed a genuine defect. The action condition was originally met "
    "only by an explicit verb ('click', 'visit') or a visible URL; every one of the 865 "
    "AI messages contains a literal URL, but only about two-fifths of the real-world "
    "messages do, because converting an HTML email to text discards the anchor's href "
    "and leaves only its label ('UPDATE MY DETAILS'). The rule was thus near-trivially "
    "satisfied for one class and structurally unsatisfiable for much of the other, and "
    "roughly 60% of the sampled excluded real-world messages proved to be genuine "
    "credential lures — a source-dependent selection effect that would have biased the "
    "real-world class towards messages that happened to spell out their call to action. "
    "The screen was corrected by adding an imperative call-to-action arm covering button "
    "text, folding invisible characters and Cyrillic homoglyphs (used to break up "
    "keywords) before matching, and matching on word stems.")
para(
    "Re-auditing the corrected screen gave a retention precision of "
    f"{audit['by_source']['nazario']['precision_of_retained']:.2f} for the real-world "
    f"class and {audit['by_source']['ai']['precision_of_retained']:.2f} for the AI "
    "class, and, weighting strata by population, estimated recalls of "
    f"{audit['by_source']['nazario']['estimated_recall']:.2f} and "
    f"{audit['by_source']['ai']['estimated_recall']:.2f}. Residual false positives are "
    "document-signature and payment-advice lures that mention an account term; residual "
    "false negatives are mailbox-recovery notices phrased unusually. The screen is "
    "therefore accurate but imperfect, and still marginally stricter on the real-world "
    "class — a limitation carried into the interpretation rather than a claim of a "
    "clean filter. The episode is worth reporting in itself: the original rule looked "
    "reproducible and defensible on paper, and only the audit revealed it was biased.")

h2("5.3 Preventing train/test leakage")
para(
    "With repeated templates in the data, the ordinary risk of leakage — the same email "
    "appearing in both training and test — was acute, and it would inflate the metrics "
    "silently. The primary guard is exact-duplicate removal before any split: a message "
    "is hashed on its lower-cased alphanumeric characters, so copies that differ only in "
    "whitespace, case or punctuation are collapsed and cannot straddle the partitions. "
    "This catches reformatted duplicates but, being an exact-after-normalisation match, "
    "it does not by itself catch structurally near-identical template variants that "
    "differ in wording; that residual risk is quantified separately by the "
    "campaign-level grouped-split evaluation in Section 6.5. The split is stratified by "
    f"class under a single fixed seed (42), yielding {n_train} training, {n_val} "
    f"validation and {n_test} test emails ({n_test // 2} per class in the test "
    "partition). The validation partition was created to provide early-stopping data "
    "for the transformer baseline; because that baseline was cut under the proposal's "
    "contingency (Section 5.5), it is not used in any reported analysis, and "
    "hyperparameters are selected by five-fold cross-validation inside the training "
    "partition only. The test partition was excluded from feature fitting, model "
    "training and hyperparameter selection, and was used only for final evaluation and "
    "the stated sensitivity analyses; the whole run is deterministic.")

h2("5.4 Feature-extraction edge cases")
para(
    "Extracting 95 numeric features from arbitrary email text surfaced a series of "
    "edge cases that each required a defensive decision. Empty or near-empty bodies "
    "produce zero-denominator ratios, so every rate is computed through a safe-division "
    "helper that returns zero rather than raising. Vocabulary-richness statistics such "
    "as Honoré's R are undefined when every word is unique or the text is a single "
    "token, so guarded formulae fall back to zero in those degenerate cases. Syllable "
    "counting for the readability indices uses a vowel-group approximation, which is "
    "accurate enough for a relative comparison between classes without importing a "
    "pronunciation dictionary. Keeping the extractor dependent only on the standard "
    "library and numpy was a deliberate constraint: it let the same code be unit-tested "
    "in isolation and guarantees that every feature can be explained in plain terms.")

h2("5.5 Hyperparameter search under modest compute")
para(
    "The Random Forest and XGBoost models are tuned by five-fold cross-validated grid "
    "search on the training data only, optimising F1. The grids were kept deliberately "
    "compact — a few hundred candidate fits rather than thousands — because the "
    "experiment had to run on a laptop, and because the stylometric signal, as it "
    "turned out, is strong enough that the models are not sensitive to fine tuning. "
    "The selected configurations were a "
    f"{results['random_forest']['best_params'].get('n_estimators','?')}-tree forest of "
    f"depth {results['random_forest']['best_params'].get('max_depth','?')} and an "
    "XGBoost model of depth "
    f"{results['xgboost']['best_params'].get('max_depth','?')} with "
    f"{results['xgboost']['best_params'].get('n_estimators','?')} rounds. A late "
    "obstacle was environmental rather than statistical: XGBoost failed to load on "
    "macOS until the OpenMP runtime it links against was installed — an environmental "
    "dependency issue that justifies pinning dependencies in requirements.txt.")

h2("5.6 Cutting the transformer baseline")
para(
    "The proposal listed a fine-tuned transformer as a Should-have comparison baseline "
    "and named it as the first item to cut if compute ran short. That contingency was "
    "exercised. Once the stylometric models were returning very high F1 scores with "
    "tight confidence intervals, a transformer baseline would have added considerable "
    "compute and complexity without changing the answer to the research question, "
    "which concerns stylometric features specifically. Cutting it was a scope decision "
    "consistent with the proposal's stated contingency plan, and its implications are "
    "discussed in Chapter 7.")

h2("5.7 Testing and sanity baselines")
para(
    "The scientific core is covered by unit tests exercising the feature extractor and "
    "the metric functions, including the degenerate inputs above; all tests pass. "
    "Beyond unit testing, the design guards against the classic failure mode of "
    "reporting an impressive number that is really an artefact. A neutral synthetic "
    "smoke-test dataset exists purely to prove the pipeline runs end-to-end and is "
    "never reported as a result. More importantly, the very high headline numbers "
    "prompted a deliberate scepticism that is addressed quantitatively in Sections 6.4 "
    "and 6.5 through a structural-cue ablation and a campaign-level grouped split, so "
    "the evaluation tests whether the result is genuine rather than merely asserting it.")
page_break()

# ============================ 6. EVALUATION ============================
h1("6. Evaluation")

h2("6.1 Headline results")
para(
    "All three classifiers were evaluated on the untouched held-out test set of "
    f"{n_test} emails. Every model exceeded the F1 = 0.90 target pre-specified in the "
    "proposal by a wide margin (Table 6.1). The logistic-regression baseline achieved "
    f"an F1 of {f('logreg','f1'):.3f} (95% bootstrap CI "
    f"[{f('logreg','f1_ci_low'):.3f}, {f('logreg','f1_ci_high'):.3f}], ROC-AUC "
    f"{f('logreg','roc_auc'):.3f}), with XGBoost at {f('xgboost','f1'):.3f} and the "
    f"Random Forest at {f('random_forest','f1'):.3f}. That a linear model performs at "
    "least as well as the tree ensembles is informative in itself: it indicates the two "
    "classes are close to linearly separable in this feature space, so the "
    "discriminative signal is broad and additive rather than concealed in complex "
    "feature interactions. Scores this high on a single split should nonetheless be "
    "read with caution rather than celebration, because one split cannot distinguish a "
    "genuinely strong signal from one inflated by repeated campaign templates or by "
    "corpus-specific cues. Sections 6.4 and 6.5 therefore subject the result to two "
    "stress tests, and the grouped-split result in Table 6.4 — not this point estimate "
    "— should be taken as the headline.")
table(
    ["Model", "Precision", "Recall", "F1", "F1 95% CI", "ROC-AUC", "Accuracy"],
    [[m.replace("logreg", "Logistic regression").replace("random_forest", "Random Forest").replace("xgboost", "XGBoost"),
      f"{f(m,'precision'):.3f}", f"{f(m,'recall'):.3f}", f"{f(m,'f1'):.3f}",
      f"[{f(m,'f1_ci_low'):.3f}, {f(m,'f1_ci_high'):.3f}]",
      f"{f(m,'roc_auc'):.3f}", f"{f(m,'accuracy'):.3f}"]
     for m in ["logreg", "random_forest", "xgboost"]],
    widths=[1.7, 0.85, 0.7, 0.6, 1.35, 0.85, 0.85])
caption("Table 6.1 — Random-split test-set performance for the positive "
        "(AI-generated) class. Interpret alongside the grouped split in Table 6.4.")

h2("6.2 Confusion matrices")
para(
    "The confusion matrices confirm that errors are few and of both kinds. The "
    f"logistic-regression model misclassified {cm_lr[0][1] + cm_lr[1][0]} of the "
    f"{sum(cm_lr[0]) + sum(cm_lr[1])} test emails, flagging {cm_lr[0][1]} real-world "
    f"messages as AI-generated while missing {'none' if cm_lr[1][0] == 0 else cm_lr[1][0]} "
    f"of the AI-generated messages (Figure 6.1); the Random Forest traded recall for "
    f"precision, missing {cm_rf[1][0]} AI-generated messages and raising "
    f"{cm_rf[0][1]} false alarm. For a defensive filter that trade-off is the "
    "operative choice: the operating point can be shifted along the ROC curve to "
    "favour precision or recall as deployment requires.")
figure(os.path.join(RES, "confusion_logreg.png"), 3.1,
       "Figure 6.1 — Confusion matrix for the logistic-regression model on the held-out "
       "test set. Rows are the true class, columns the prediction.")

h2("6.3 Which stylistic markers reveal machine authorship")
para(
    "The second half of the aim asks which markers carry the signal. Two importance "
    "methods were computed on the Random Forest — impurity-based and permutation-based — "
    "and per-class means (Table 6.2) give the direction of each effect. The clearest "
    "and most interpretable separator is readability: AI-generated emails are markedly "
    "harder to read than real-world phishing, with a mean Flesch Reading Ease of "
    f"{cm['flesch_reading_ease']['llm_mean']:.1f} against "
    f"{cm['flesch_reading_ease']['human_mean']:.1f}, and a correspondingly higher "
    f"Flesch–Kincaid grade ({cm['flesch_kincaid_grade']['llm_mean']:.1f} versus "
    f"{cm['flesch_kincaid_grade']['human_mean']:.1f}). The AI prose is lexically "
    f"heavier (mean word length {cm['avg_word_length']['llm_mean']:.2f} versus "
    f"{cm['avg_word_length']['human_mean']:.2f} characters) and leans more on urgency "
    f"vocabulary ({cm['urgency_word_ratio']['llm_mean']:.3f} versus "
    f"{cm['urgency_word_ratio']['human_mean']:.3f}), whereas real-world phishing carries "
    f"more digits ({cm['digit_ratio']['human_mean']:.3f} versus "
    f"{cm['digit_ratio']['llm_mean']:.3f}) and heavier terminal punctuation. These "
    "directions are coherent with the theory in Chapter 2: the fluency and "
    "register-uniformity that make LLM output persuasive also make it measurably "
    "regular, and that regularity is what the classifier exploits.")
table(
    ["Feature", "Real-world mean", "AI mean", "Reading of the signal"],
    [["Flesch Reading Ease", f"{cm['flesch_reading_ease']['human_mean']:.1f}",
      f"{cm['flesch_reading_ease']['llm_mean']:.1f}", "AI text is harder to read"],
     ["Flesch–Kincaid grade", f"{cm['flesch_kincaid_grade']['human_mean']:.1f}",
      f"{cm['flesch_kincaid_grade']['llm_mean']:.1f}", "AI writes at a higher grade"],
     ["Avg. word length", f"{cm['avg_word_length']['human_mean']:.2f}",
      f"{cm['avg_word_length']['llm_mean']:.2f}", "AI uses longer words"],
     ["Urgency-word ratio", f"{cm['urgency_word_ratio']['human_mean']:.3f}",
      f"{cm['urgency_word_ratio']['llm_mean']:.3f}", "AI leans on urgency vocabulary"],
     ["Digit ratio", f"{cm['digit_ratio']['human_mean']:.3f}",
      f"{cm['digit_ratio']['llm_mean']:.3f}", "Real-world phishing carries more digits"],
     ["Word count", f"{cm['word_count']['human_mean']:.0f}",
      f"{cm['word_count']['llm_mean']:.0f}", "AI messages are shorter"]],
    widths=[1.7, 1.25, 0.85, 2.5])
caption("Table 6.2 — Direction of the strongest signals: per-class feature means over "
        "the full corpus.")
para(
    "The two importance methods must be reported together rather than merged, because "
    "they do not agree closely, and saying so is more honest than presenting a single "
    "ranking as settled. Impurity importance is dominated by the readability and "
    "word-length cluster (average word length, Flesch–Kincaid grade, Flesch Reading "
    "Ease, syllables per word), whereas permutation importance promotes function-word "
    "frequencies and structural counts (the relative frequency of 'the', "
    "email-address count, 'please', Honoré's R). Only "
    f"{agree['features_in_both_top20']} features appear in both top-20 lists, and "
    "across those the Spearman rank correlation is "
    f"{agree['spearman_rho']:.2f} (p = {agree['p_value']:.2f}) — no reliable agreement "
    "on ordering. The divergence is explicable: impurity importance is computed on "
    "training data and favours continuous, high-cardinality features such as the "
    "readability indices, whereas permutation importance measures the test-set cost of "
    "destroying a feature and so rewards what the fitted model actually relies on. "
    "Together they support a weaker but defensible claim — that both a readability and "
    "length cluster and a function-word cluster carry discriminative information — "
    "rather than one definitive ranking (Table 6.3).")
table(
    ["Rank", "Impurity-based (RF)", "Permutation-based (RF)"],
    [[str(i + 1), supp["rf_impurity_top20"][i]["feature"],
      supp["rf_permutation_top20"][i]["feature"]] for i in range(8)],
    widths=[0.7, 2.9, 2.9])
caption("Table 6.3 — Top eight features by each importance method. The two methods "
        f"share only {agree['features_in_both_top20']} features across their full "
        f"top-20 lists (Spearman ρ = {agree['spearman_rho']:.2f}), so both are "
        "reported rather than combined.")
figure(os.path.join(RES, "feature_importance.png"), 6.2,
       "Figure 6.2 — Random-Forest impurity importance (top 20 features).")
figure(os.path.join(RES, "fig_readability.png"), 5.0,
       "Figure 6.3 — Distribution of Flesch Reading Ease by class. The real-world "
       "distribution sits markedly higher (easier to read); dashed lines mark the "
       "class means.")

h2("6.4 Structural-cue ablation")
para(
    "One pair of features demanded scrutiny before the result could be trusted. The "
    "email-address count is almost perfectly separating — real-world phishing from the "
    f"Nazario mailboxes contains real addresses (mean "
    f"{cm['email_addr_count']['human_mean']:.2f} per message) while the AI corpus, "
    f"generated without real recipients, contains essentially none "
    f"({cm['email_addr_count']['llm_mean']:.2f}) — and URL count differs similarly. "
    "These are properties of how each corpus was collected rather than of writing "
    "style. To test their influence both were removed and the models retrained. "
    f"Performance changed little: logistic-regression F1 moved from {f('logreg','f1'):.3f} "
    f"to {abl['logreg']['f1']:.3f}, XGBoost from {f('xgboost','f1'):.3f} to "
    f"{abl['xgboost']['f1']:.3f} and the Random Forest from {f('random_forest','f1'):.3f} "
    f"to {abl['random_forest']['f1']:.3f}, with overlapping confidence intervals in "
    "every case. This indicates the result is not driven solely by those two obvious "
    "structural cues. It "
    "does not, however, establish the absence of all corpus-source confounding: "
    "differences in collection process, prompt-template regularities, greetings and "
    "sign-offs, length or topic distributions could still contribute, and would require "
    "cross-source or matched-topic evaluation to exclude fully.")

h2("6.5 Campaign-level (grouped) split and 2022-only sensitivity")
para(
    "Because phishing campaigns repeat a template with small edits, a random split can "
    "place near-identical variants in both training and test and inflate performance. "
    "To bound this, all "
    f"{gsc['n_emails']} emails were clustered by textual similarity (character 5-gram "
    f"shingles, Jaccard ≥ {gsc['jaccard_threshold']}), yielding {gsc['n_clusters']} "
    f"clusters ({gsc['singletons']} singletons, largest cluster {gsc['largest_cluster']} "
    "emails), and the split was redrawn so that every cluster falls wholly within one "
    "partition — so no messages belonging to the same detected similarity cluster were "
    "permitted to straddle the training and test partitions. On the resulting "
    f"{gs['n_train']}/{gs['n_test']} split ({gs['test_positive']} AI-generated and "
    f"{gs['test_negative']} real-world test messages) the logistic-regression F1 was "
    f"{gs['logreg']['f1']:.3f} (95% CI [{gs['logreg']['f1_ci'][0]:.3f}, "
    f"{gs['logreg']['f1_ci'][1]:.3f}]), XGBoost {gs['xgboost']['f1']:.3f} and the "
    f"Random Forest {gs['random_forest']['f1']:.3f} (Table 6.4). Performance did not "
    "fall relative to the random split; if anything it was marginally higher. Had "
    "template leakage been inflating the random-split figures the expected pattern was "
    "a clear drop, so its absence is evidence that leakage between near-identical "
    "campaign variants is not what the classifiers were exploiting. This is the key "
    "robustness result, but it should not be over-read: a single grouped partition has "
    "its own composition and can happen to be slightly easier, so the defensible "
    "conclusion is that the signal is insensitive to campaign-level leakage, not that "
    "grouping improves detection. Repeated grouped splits would sharpen the estimate.")
def _fci(d, model):
    """F1 with its bootstrap interval, for the robustness table."""
    m = d[model]
    lo, hi = (m["f1_ci"] if "f1_ci" in m else [m["f1_ci_low"], m["f1_ci_high"]])
    return f"{m['f1']:.3f}\n[{lo:.3f}, {hi:.3f}]"


table(
    ["Evaluation", "Test n", "Pos.", "Logistic F1", "Random Forest F1", "XGBoost F1"],
    [["Random split (Table 6.1)", str(n_test), str(n_test // 2),
      _fci(results, "logreg"), _fci(results, "random_forest"), _fci(results, "xgboost")],
     ["Structural-cue ablation", str(abl["n_test"]), str(n_test // 2),
      _fci(abl, "logreg"), _fci(abl, "random_forest"), _fci(abl, "xgboost")],
     ["Campaign-level grouped split", str(gs["n_test"]), str(gs["test_positive"]),
      _fci(gs, "logreg"), _fci(gs, "random_forest"), _fci(gs, "xgboost")],
     [f"2022-only subset ({s2022['n_per_class']}/class)", str(s2022["n_test"]),
      str(s2022["test_positive"]),
      _fci(s2022, "logreg"), _fci(s2022, "random_forest"), _fci(s2022, "xgboost")]],
    widths=[1.55, 0.5, 0.42, 1.25, 1.3, 1.25])
caption("Table 6.4 — F1 with 95% bootstrap intervals under robustness and sensitivity "
        "conditions, for all three classifiers. The signal persists under grouped "
        "splitting and on the 2022-only sensitivity subset.")
para(
    "One caveat applies to Table 6.4. Where a model classified every test message "
    "correctly the bootstrap interval collapses to [1.000, 1.000]; that is an artefact "
    "of resampling an error-free prediction set, not a claim that generalisation is "
    "known exactly. The interval captures sampling variation in the test predictions "
    "alone, not uncertainty from model fitting, corpus construction or choice of "
    "partition, which would require repeated group-aware cross-validation to quantify.",
    after=6)
para(
    "A second sensitivity analysis addresses the authorship uncertainty of the "
    "real-world class. Because the 2023–24 archives overlap heavy public LLM use, the "
    "real-world class was restricted to the 2022 archive — the least likely to contain "
    f"AI-assisted text — and rebalanced ({s2022['n_per_class']} per class, "
    f"{s2022['n_test']} held-out test messages). The signal survived (logistic F1 "
    f"{s2022['logreg']['f1']:.3f}, Random Forest {s2022['random_forest']['f1']:.3f}, "
    f"XGBoost {s2022['xgboost']['f1']:.3f}), though the small sample gives wide "
    "confidence intervals (Table 6.4). Two caveats apply. This cannot prove human "
    "authorship; it shows only that the result does not depend on the most "
    "exposure-prone years. And restricting the real-world class to 2022 while the AI "
    "class remains a 2024 dataset increases the temporal mismatch between the classes, "
    "so this is an authorship-contamination sensitivity check rather than an "
    "independent confirmation free of chronological confounding.")

h2("6.6 Revisiting the aim, and what the evaluation cannot show")
para(
    "The aim was to determine the extent to which stylometric features distinguish "
    "AI-generated from real-world phishing and to identify the most discriminative "
    "features. Both halves are answered: the separation is strong and survives "
    "structural-cue ablation, campaign-level grouped splitting and a 2022-only subset, "
    f"with F1 between {min(abl['logreg']['f1'], abl['random_forest']['f1'], abl['xgboost']['f1'], s2022['logreg']['f1']):.3f} "
    f"and {max(gs['logreg']['f1'], gs['xgboost']['f1']):.3f} across conditions; and the signal is carried by "
    "readability, word length, punctuation and urgency vocabulary. The target of F1 = "
    "0.90 was met under every condition tested. Three boundaries must nonetheless be "
    "stated plainly. The AI class comes from a single published generation dataset, so "
    "these metrics measure the detectability of that dataset's output, not of AI "
    "phishing in general. The real-world class is not verified as human-authored. And "
    "the evaluation is static: it does not model an adversary who paraphrases or "
    "'humanises' machine output to defeat a stylometric detector, which the literature "
    "shows is effective against related methods (Krishna et al., 2023). High held-out "
    "performance is necessary but not sufficient for a deployable detector, and "
    "Chapter 7 treats these boundaries as the agenda for future work.")
page_break()

# ============================ 7. DISCUSSION ============================
h1("7. Discussion")

h2("7.1 Does writing style betray machine authorship?")
para(
    "The direct answer to the research question is yes, within the scope studied. With "
    "no access to the generating model, a transparent classifier over 95 interpretable "
    "stylometric features separates known AI-generated from real-world phishing with an "
    f"F1 that stays high even under the stricter tests — {gs['logreg']['f1']:.3f} for "
    "the logistic model under a campaign-level grouped split, and "
    f"{s2022['logreg']['f1']:.3f} on the 2022-only subset. The "
    "solution the project set out to build — a lightweight, transparent, content-only "
    "detection signal — works on these two corpora. It is not a black box: it can be "
    "explained in a sentence, namely that the AI-generated phishing in this dataset "
    "reads as more formal, more lexically heavy and more uniform than the real-world "
    "variety, and that this regularity is measurable. The strength of the raw "
    "random-split score (a perfect F1 for the tree models) is best read as a sign that "
    "the two particular sources are highly separable once restricted to credential "
    "harvesting, rather than as a claim about AI phishing in general.")

h2("7.2 Relationship to the literature")
para(
    "The finding replicates and extends Eze and Shamir (2024), who used the stylistic "
    "UDAT method to separate AI-generated phishing from human email and ranked several "
    "informative descriptors; this project reproduces the core result on their AI "
    "corpus against a real-world phishing class, narrows it to a single attack type, and "
    "adds ranked interpretation, uncertainty estimates and leakage controls. It is also "
    "consistent with the classical stylometry of Mosteller and Wallace (1964) and "
    "Stamatatos (2009): the same premise that unconscious stylistic habits betray "
    "authorship transfers from the human-versus-human setting to human-versus-machine. "
    "At the same time the work inherits, rather than escapes, the central caution of "
    "the machine-text-detection literature. Mitchell et al. (2023) and especially "
    "Krishna et al. (2023) show that stylistic and statistical detectors are brittle "
    "under paraphrasing; there is no reason to think a stylometric detector is immune, "
    "and the readability signal that carries most of the discrimination here is exactly "
    "the kind of surface property a 'humanising' paraphrase would target. The honest "
    "reading is that stylometry is a strong signal against un-adapted LLM phishing and "
    "an unproven one against an adaptive adversary.")

h2("7.3 Reflection on supervisor feedback and the project journey")
para(
    "The project changed in several ways since the proposal, in response to supervisory "
    "guidance to keep the study defensible and ethically clean. The proposal envisaged "
    "generating the AI phishing locally from a chosen model under controlled prompts. To "
    "remove the dual-use hazard of producing fresh attack content, the design switched "
    "to the existing published Eze and Shamir (2024) corpus. This improved both the "
    "ethics and the reproducibility — anyone can obtain the same data and reproduce the "
    "numbers — at the acknowledged cost of tying the AI class to one published "
    "generation dataset. A second change, prompted by the reasonable suspicion that "
    "near-perfect scores might reflect corpus artefacts rather than style, was to add "
    "the structural-cue ablation (Section 6.4) and, in response to review, the "
    "campaign-level grouped split and the 2022-only sensitivity analysis (Sections 6.5); "
    "addressing that suspicion quantitatively rather than rhetorically materially "
    "strengthened the evaluation, and the honest outcome — a small drop under grouped "
    "splitting — is reported rather than hidden. A third change was relabelling the "
    "classes from 'human-written' to 'real-world (Nazario) phishing' once it was clear "
    "the source guarantees phishing classification but not verified human authorship. A "
    "fourth, smaller adjustment was cutting the transformer baseline in line with the "
    "proposal's stated contingency, once it was clear it would not change the answer.")

h2("7.4 Progress against the project plan")
para(
    "Measured against the two-week schedule and milestones in the proposal, the shape of "
    "the work changed more than its total effort. Corpus construction, planned as an "
    "early and largely mechanical task, became the critical path and consumed "
    "substantially more time than allowed: parsing two structurally incompatible "
    "sources, adding a credential-harvesting rule, auditing it and then correcting it "
    "absorbed effort originally reserved for the transformer baseline, which was cut "
    "under the proposal's contingency. Evaluation expanded well beyond the plan: "
    "bootstrap intervals, a structural-cue ablation, campaign-level clustering with a "
    "grouped split and a 2022-only sensitivity analysis were added in response to the "
    "risk that a high headline score reflected corpus artefacts rather than style. The "
    "net effect was to shift effort from model variety towards data validity, which was "
    "the right trade for a study whose central threat was never insufficient accuracy "
    "but insufficient confidence that the accuracy meant what it appeared to. Weekly "
    "supervisor checkpoints were retained and the submission milestone was met.")

h2("7.5 Does the artefact solve the stated problem?")
para(
    "The problem defined in Chapter 1 was the detection of machine-authored phishing "
    "through its writing style. Against that problem the artefact is a success on its "
    "own terms: it demonstrates that a strong, interpretable stylistic signal separates "
    "the two corpora, that the signal is not driven solely by the obvious structural "
    "cues, and that it survives campaign-level leakage control. It does not, and did not "
    "claim to, deliver a production-ready detector robust to an adaptive attacker across "
    "every generator and language, nor does it exclude all possible source-level "
    "confounding between the two corpora. The contribution is a clear, reproducible, "
    "evidence-backed characterisation of a precisely scoped signal — which is what a "
    "project of this size should aim to provide — together with a pipeline others can "
    "extend toward the harder, open version of the problem.")
page_break()

# ============================ 8. CONCLUSION ============================
h1("8. Conclusion")
para(
    "This study found that a transparent set of stylometric and lexical features "
    "distinguished known AI-generated messages from real-world Nazario phishing with an "
    f"F1 of approximately {gs['logreg']['f1']:.2f} on a campaign-level held-out split "
    "(and higher still on a simple random split). Readability, word length, punctuation "
    "and selected lexical features such as urgency vocabulary contributed most strongly "
    "to the classification. Removing the URL and email-address counts produced little "
    "loss in performance, showing that the result was not driven solely by those two "
    "structural cues, and the signal persisted under campaign-level grouped splitting "
    "and on the 2022-only sensitivity subset. Nevertheless, the findings remain "
    "specific to two source corpora and one published AI-generation dataset; verified "
    "human authorship, cross-generator generalisation, campaign-level leakage beyond "
    "the clustering used here, and broader source confounding all require further "
    "investigation. The practical significance is that stylometry offers a lightweight, "
    "transparent detection signal that does not require access to the generating model "
    "at inference, and that could complement "
    "existing content- and header-based phishing filters.")
para(
    "The scope that makes the answer defensible also bounds it, and several lines of "
    "future work follow directly. The most important is cross-generator robustness: "
    "repeating the study with phishing from several models and services would test "
    "whether the readability signature generalises or is specific to this dataset. A "
    "second is stronger confound control — matched-topic or matched-length sampling, "
    "removal of greetings and organisation names, and evaluation on a fully independent "
    "corpus. A third is adversarial robustness — measuring how far deliberate "
    "paraphrasing or 'humanisation' degrades the detector, following Krishna et al. "
    "(2023). A fourth is extension beyond English-language credential-harvesting email. "
    "Addressing these would move the work from a well-characterised result on two "
    "corpora towards a detector fit for the adaptive, multilingual threat that "
    "AI-generated phishing is becoming.")
page_break()

# ============================ REFERENCES ============================
h1("References")
refs = [
 "Abu-Nimeh, S., Nappa, D., Wang, X. and Nair, S. (2007) 'A comparison of machine learning techniques for phishing detection', Proceedings of the Anti-Phishing Working Group eCrime Researchers Summit, pp. 60-69.",
 "Anti-Phishing Working Group (2025) Phishing Activity Trends Report. Cambridge, MA: Anti-Phishing Working Group. Available at: https://apwg.org/trendsreports/ (Accessed: 25 July 2026).",
 "Computer Misuse Act 1990, c. 18. London: HMSO.",
 "Data Protection Act 2018, c. 12. London: The Stationery Office.",
 "Devlin, J., Chang, M.-W., Lee, K. and Toutanova, K. (2019) 'BERT: pre-training of deep bidirectional transformers for language understanding', Proceedings of NAACL-HLT 2019, pp. 4171-4186.",
 "Eze, C.S. and Shamir, L. (2024) 'Analysis and prevention of AI-based phishing email attacks', Electronics, 13(10), 1839.",
 "Fette, I., Sadeh, N. and Tomasic, A. (2007) 'Learning to detect phishing emails', Proceedings of the 16th International Conference on World Wide Web (WWW '07), pp. 649-656.",
 "Gehrmann, S., Strobelt, H. and Rush, A.M. (2019) 'GLTR: statistical detection and visualization of generated text', Proceedings of ACL 2019: System Demonstrations, pp. 111-116.",
 "Heiding, F., Schneier, B., Vishwanath, A., Bernstein, J. and Park, P.S. (2024) 'Devising and detecting phishing emails using large language models', IEEE Access, 12, pp. 42131-42146.",
 "International Organization for Standardization (2022) ISO/IEC 27001:2022 Information security, cybersecurity and privacy protection - Information security management systems - Requirements. Geneva: ISO.",
 "Kirchenbauer, J., Geiping, J., Wen, Y., Katz, J., Miers, I. and Goldstein, T. (2023) 'A watermark for large language models', Proceedings of the 40th International Conference on Machine Learning (ICML), PMLR 202, pp. 17061-17084.",
 "Koide, T., Fukushi, N., Nakano, H. and Chiba, D. (2024) ChatSpamDetector: Leveraging Large Language Models for Effective Phishing Email Detection. arXiv:2402.18093.",
 "Krishna, K., Song, Y., Karpinska, M., Wieting, J. and Iyyer, M. (2023) 'Paraphrasing evades detectors of AI-generated text, but retrieval is an effective defense', Advances in Neural Information Processing Systems (NeurIPS) 36.",
 "Mitchell, E., Lee, Y., Khazatsky, A., Manning, C.D. and Finn, C. (2023) 'DetectGPT: zero-shot machine-generated text detection using probability curvature', Proceedings of the 40th International Conference on Machine Learning (ICML), PMLR 202, pp. 24950-24962.",
 "Mosteller, F. and Wallace, D.L. (1964) Inference and Disputed Authorship: The Federalist. Reading, MA: Addison-Wesley.",
 "Nazario, J. (2005-2025) Phishing Corpus [dataset]. Licence: CC BY 4.0. Available at: https://monkey.org/~jose/phishing/ (Accessed: 22 July 2026).",
 "National Cyber Security Centre (2024) Phishing attacks: defending your organisation. London: NCSC. Available at: https://www.ncsc.gov.uk/guidance/phishing (Accessed: 25 July 2026).",
 "Solaiman, I., Brundage, M., Clark, J., Askell, A., Herbert-Voss, A., Wu, J., Radford, A. et al. (2019) Release Strategies and the Social Impacts of Language Models. arXiv:1908.09203.",
 "Stamatatos, E. (2009) 'A survey of modern authorship attribution methods', Journal of the American Society for Information Science and Technology, 60(3), pp. 538-556.",
 "UK General Data Protection Regulation (2021). Retained Regulation (EU) 2016/679.",
 "Verizon (2025) 2025 Data Breach Investigations Report. New York: Verizon Business. Available at: https://www.verizon.com/business/resources/reports/dbir/ (Accessed: 25 July 2026).",
 "Zellers, R., Holtzman, A., Rashkin, H., Bisk, Y., Farhadi, A., Roesner, F. and Choi, Y. (2019) 'Defending against neural fake news', Advances in Neural Information Processing Systems (NeurIPS) 32.",
]
for rtext in refs:
    p = doc.add_paragraph(rtext)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
page_break()

# ============================ APPENDICES ============================
h1("Appendix A: Artefact and source code")
_ap = doc.add_paragraph()
_ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
_ap.add_run(
    "The complete artefact — feature extractor, data loader, training, evaluation and "
    "interpretation modules, unit tests, and the corpus-construction and analysis "
    "scripts that produced every table and figure in this report — is version-controlled "
    "in the MS-Artifacts/ folder of the repository at ")
add_hyperlink(_ap, "https://github.com/Mivics1/MS-Assignments",
              "https://github.com/Mivics1/MS-Assignments")
_ap.add_run(
    f". The exact state used for this report is release {RELEASE_TAG}, commit "
    f"{COMMIT_HASH}. The full experiment is reproduced from the raw sources by running, "
    "from the MS-Artifacts directory and under the default seed of 42:")
for step in [
    "python scripts/build_corpus.py  — rebuild data/emails.csv from the raw sources",
    "python -m src.pipeline --data data/emails.csv --out results/  — train and evaluate",
    "python scripts/supplementary_analysis.py  — class means, importance, ablation",
    "python scripts/grouped_split_eval.py  — grouped split and 2022 sensitivity",
    "python scripts/audit_screen_sample.py and audit_screen_score.py  — screening audit",
    "python scripts/make_figures.py  — architecture and readability figures",
    "python scripts/build_final_dissertation.py  — regenerate this document",
]:
    numbered(step)
para(
    "Raw phishing text is excluded from the public repository for the ethical reasons "
    "set out in Section 3.6; only aggregate statistics, derived features and dataset "
    "checksums are published. The corpus-construction script documents exactly how to "
    "rebuild data/emails.csv from the two cited public sources, and the checksums in "
    "Table B.3 allow a reader to confirm they obtained identical raw inputs.", after=6)

h1("Appendix B: Corpus construction and provenance")
para(
    "The tables below record the corpus construction and source provenance as produced "
    "by the build script, for full reproducibility. The credential-harvesting screen "
    "and its exclusion counts are reported so the scope is evidenced rather than "
    "asserted.", after=6)
prov = corpus["provenance"]
fnl = corpus["funnel"]
nz, az = fnl["nazario"], fnl["eze_shamir"]
table(
    ["Stage", "Real-world (Nazario)", "AI (Eze & Shamir)"],
    [["Raw messages parsed", str(nz["parsed"]), str(az["parsed"])],
     ["Folder-internal pseudo-messages removed",
      str(nz["excluded_folder_internal"]), str(az["excluded_folder_internal"])],
     ["Removed by length filter",
      str(nz["excluded_length_filter"]), str(az["excluded_length_filter"])],
     ["Removed as non-English",
      str(nz["excluded_non_english"]), str(az["excluded_non_english"])],
     ["Excluded by credential-harvesting screen",
      str(nz["excluded_not_credential_harvesting"]),
      str(az["excluded_not_credential_harvesting"])],
     ["Retained before duplicate removal",
      str(nz["retained_before_dedupe"]), str(az["retained_before_dedupe"])],
     ["Exact duplicates removed",
      str(nz["exact_duplicates_removed"]), str(az["exact_duplicates_removed"])],
     ["Remaining after de-duplication",
      str(nz["after_dedupe"]), str(az["after_dedupe"])],
     ["Removed by balancing down-sample",
      str(nz["removed_by_downsampling"]), str(az["removed_by_downsampling"])],
     ["Final class size", str(nz["final"]), str(az["final"])]],
    widths=[3.1, 1.6, 1.6])
caption(f"Table B.1 — Corpus construction and screening funnel; every stage is "
        f"reported so the counts reconcile to the final {corpus['total']}-message "
        "corpus.")
table(
    ["Item", "Value"],
    [["Real-world (class 0) source", "Nazario phishing corpus 2022-2024 (CC BY 4.0)"],
     ["AI (class 1) source",
      "Eze and Shamir (2024), 865 emails via the DeepAI text-generation service"],
     ["Length filter (chars)", f"{corpus['min_chars']}-{corpus['max_chars']}"],
     ["Credential screen", corpus["credential_harvesting_screen"]["rule"]],
     ["Screen precision (real-world / AI)",
      f"{audit['by_source']['nazario']['precision_of_retained']:.2f} / "
      f"{audit['by_source']['ai']['precision_of_retained']:.2f} (audited, n=100)"],
     ["Screen estimated recall (real-world / AI)",
      f"{audit['by_source']['nazario']['estimated_recall']:.2f} / "
      f"{audit['by_source']['ai']['estimated_recall']:.2f}"],
     ["Download date", prov["download_date"]],
     ["Random seed", str(corpus["seed"])]],
    widths=[2.4, 3.9])
caption("Table B.2 — Corpus parameters, source provenance and audited screen "
        "performance.")
table(
    ["Source file", "SHA-256 (first 32 hex chars)"],
    [[k, v[:32]] for k, v in prov["sha256"].items()],
    widths=[3.0, 3.3])
caption("Table B.3 — Provenance checksums of the raw source files (full hashes in "
        "data/corpus_summary.json).")

# ---- fill List of Figures / List of Tables from collected captions ----
a = LOF_ANCHOR
for cap in FIGLIST:
    a = insert_after(a, cap)
if not FIGLIST:
    insert_after(LOF_ANCHOR, "(no figures)")
a = LOT_ANCHOR
for cap in TBLLIST:
    a = insert_after(a, cap)
if not TBLLIST:
    insert_after(LOT_ANCHOR, "(no tables)")

# ---- page numbering: prelims in lower-roman (no number on the title page),
#      body in decimal restarting at 1 from Chapter 1 ----
prelim = doc.sections[0]
set_page_numbering(prelim, "lowerRoman")
footer_page_number(prelim, blank_first=True)
set_page_numbering(BODY_SECTION, "decimal", start=1)
footer_page_number(BODY_SECTION, blank_first=False)
BODY_SECTION.footer.is_linked_to_previous = False

# ---- tell Word to refresh all fields (TOC, LoF page refs) on open ----
_settings = doc.settings.element
if _settings.find(qn("w:updateFields")) is None:
    uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
    _settings.append(uf)

doc.save(OUT)
print("saved:", OUT)
