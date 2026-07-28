"""
Validate that the finalised dissertation agrees with freshly reproduced
artefact outputs.

The dissertation is authoritative for wording, scope and interpretation; the
artefact is authoritative for every number. This script re-derives the numbers
from the artefact, recomputes the metrics from the confusion matrices, and
compares them both with the stored result files and with the tables actually
printed in the dissertation .docx.

Run:  PYTHONPATH=. .venv/bin/python scripts/validate_final_alignment.py [docx]

Exits 1 if any required check fails. Writes results/final_alignment_validation.json.

Author: Agboola Michael Daramola
Module: CO4011 Master's Project
"""

from __future__ import annotations

import csv
import datetime as _dt
import json
import os
import platform
import re
import sys

import numpy as np

RAW_TOL = 1e-6        # recomputed metrics vs stored metrics
RPT_TOL = 0.0005      # values reported to three decimal places
MEAN_TOL = 0.05       # feature means reported to 1-2 decimal places

RESULTS = "results"
DOCX_DEFAULT = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "..", "MS-Final-Dissertation", "Dissertation_Final_Validated.docx")

checks: list[dict] = []


def check(name: str, ok: bool, detail: str = "", group: str = "general"):
    checks.append({"check": name, "group": group, "pass": bool(ok), "detail": detail})
    return ok


def approx(a, b, tol):
    return abs(float(a) - float(b)) <= tol


# ---------------------------------------------------------------- corpus ----
def validate_corpus(summary):
    rows = list(csv.DictReader(open("data/emails.csv", encoding="utf-8")))
    total = len(rows)
    labels = [r["label"] for r in rows]
    n0, n1 = labels.count("0"), labels.count("1")

    check("corpus total is 982", total == 982, f"got {total}", "corpus")
    check("491 messages per class", n0 == 491 and n1 == 491,
          f"class0={n0} class1={n1}", "corpus")
    check("no empty message text",
          all(r["text"].strip() for r in rows), "", "corpus")
    check("labels are exactly {0,1}", set(labels) == {"0", "1"},
          f"found {sorted(set(labels))}", "corpus")

    from scripts.build_corpus import dedupe_key
    keys = [dedupe_key(r["text"]) for r in rows]
    dupes = len(keys) - len(set(keys))
    check("no duplicate normalised hashes remain", dupes == 0,
          f"{dupes} duplicates", "corpus")

    # Appendix B funnel reconciles at every stage
    ok = True
    detail = []
    for src, f in summary["funnel"].items():
        after_clean = (f["parsed"] - f["excluded_folder_internal"]
                       - f["excluded_length_filter"] - f["excluded_non_english"]
                       - f["excluded_not_credential_harvesting"])
        step1 = after_clean == f["retained_before_dedupe"]
        step2 = (f["retained_before_dedupe"] - f["exact_duplicates_removed"]
                 == f["after_dedupe"])
        step3 = f["after_dedupe"] - f["removed_by_downsampling"] == f["final"]
        ok &= step1 and step2 and step3
        detail.append(f"{src}: clean={step1} dedupe={step2} balance={step3}")
    check("Appendix B corpus counts reconcile", ok, "; ".join(detail), "corpus")
    check("corpus total equals sum of class finals",
          summary["total"] == summary["funnel"]["nazario"]["final"]
          + summary["funnel"]["eze_shamir"]["final"], "", "corpus")
    return rows


# ----------------------------------------------------------------- split ----
def validate_split(results):
    from src import data_loader
    df = data_loader.load_csv("data/emails.csv")
    sp = data_loader.stratified_split(df, seed=42)
    ntr, nva, nte = len(sp.y_train), len(sp.y_val), len(sp.y_test)
    check("random split is 686/148/148", (ntr, nva, nte) == (686, 148, 148),
          f"got {ntr}/{nva}/{nte}", "split")
    pos = int(np.sum(sp.y_test == 1))
    neg = int(np.sum(sp.y_test == 0))
    check("random test contains 74 per class", pos == 74 and neg == 74,
          f"pos={pos} neg={neg}", "split")
    check("partitions sum to corpus size", ntr + nva + nte == 982,
          f"{ntr + nva + nte}", "split")
    return nte


# ---------------------------------------------------------------- models ----
def metrics_from_cm(cm):
    (tn, fp), (fn, tp) = cm
    prec = tp / (tp + fp) if tp + fp else 0.0
    rec = tp / (tp + fn) if tp + fn else 0.0
    f1 = 2 * prec * rec / (prec + rec) if prec + rec else 0.0
    acc = (tp + tn) / (tp + tn + fp + fn)
    return prec, rec, f1, acc


def validate_models(results, n_test):
    lr = results["logreg"]
    check("logistic confusion matrix is TN72 FP2 FN0 TP74",
          lr["confusion_matrix"] == [[72, 2], [0, 74]],
          f"got {lr['confusion_matrix']}", "models")
    for name, r in results.items():
        cm = r["confusion_matrix"]
        total = sum(cm[0]) + sum(cm[1])
        check(f"{name}: confusion matrix sums to test n",
              total == n_test, f"{total} vs {n_test}", "models")
        prec, rec, f1, acc = metrics_from_cm(cm)
        ok = (approx(prec, r["precision"], RAW_TOL) and approx(rec, r["recall"], RAW_TOL)
              and approx(f1, r["f1"], RAW_TOL) and approx(acc, r["accuracy"], RAW_TOL))
        check(f"{name}: precision/recall/F1/accuracy recompute from matrix", ok,
              f"recomputed P={prec:.6f} R={rec:.6f} F1={f1:.6f} A={acc:.6f}", "models")
        check(f"{name}: F1 within its bootstrap interval",
              r["f1_ci_low"] - RPT_TOL <= r["f1"] <= r["f1_ci_high"] + RPT_TOL,
              f"{r['f1']:.3f} in [{r['f1_ci_low']:.3f},{r['f1_ci_high']:.3f}]", "models")


# ------------------------------------------------------------ robustness ----
def validate_robustness(supp, grouped):
    check("ablation removes only URL and email-address counts",
          sorted(supp["ablation_dropped"]) == ["email_addr_count", "url_count"],
          f"dropped {supp['ablation_dropped']}", "robustness")

    gs = grouped["grouped_split"]
    check("grouped partition sizes sum to corpus size",
          gs["n_train"] + gs["n_test"] == 982,
          f"{gs['n_train']}+{gs['n_test']}", "robustness")
    check("grouped test class counts sum to grouped test n",
          gs["test_positive"] + gs["test_negative"] == gs["n_test"],
          f"{gs['test_positive']}+{gs['test_negative']}={gs['n_test']}", "robustness")

    # clusters must not straddle the grouped train/test partitions
    from src import data_loader
    from scripts.grouped_split_eval import cluster, grouped_split
    df = data_loader.load_csv("data/emails.csv")
    cl = cluster(df["text"].tolist())
    tr, te = grouped_split(df, cl)
    crossing = set(cl[tr]) & set(cl[te])
    check("similarity clusters do not cross grouped partitions",
          len(crossing) == 0, f"{len(crossing)} clusters straddle", "robustness")
    check("grouped split reproduces stored partition sizes",
          len(tr) == gs["n_train"] and len(te) == gs["n_test"],
          f"recomputed {len(tr)}/{len(te)}", "robustness")

    s = grouped["sensitivity_2022_only"]
    check("2022-only subset is balanced",
          s["test_positive"] == s["test_negative"],
          f"pos={s['test_positive']} neg={s['test_negative']}", "robustness")
    # The 2022 subset is split by the same stratified splitter, so it also
    # yields a validation partition; reproduce the split to confirm the stored
    # train/test sizes rather than assuming train+test spans the subset.
    import pandas as pd
    src = pd.read_csv("data/emails.csv")
    naz22 = src[src["source"] == "phishing-2022"]
    ai = src[src["label"] == 1]
    m = min(len(naz22), len(ai))
    sub = pd.concat([naz22.sample(m, random_state=42),
                     ai.sample(m, random_state=42)]).sample(
                         frac=1, random_state=42).reset_index(drop=True)
    sp = data_loader.stratified_split(sub, seed=42)
    check("2022-only per-class count reproduces", m == s["n_per_class"],
          f"recomputed {m} vs stored {s['n_per_class']}", "robustness")
    check("2022-only train/test sizes reproduce",
          len(sp.y_train) == s["n_train"] and len(sp.y_test) == s["n_test"],
          f"recomputed {len(sp.y_train)}/{len(sp.y_test)} vs stored "
          f"{s['n_train']}/{s['n_test']}", "robustness")
    check("2022-only partitions sum to the subset size",
          len(sp.y_train) + len(sp.y_val) + len(sp.y_test) == 2 * m,
          f"{len(sp.y_train)}+{len(sp.y_val)}+{len(sp.y_test)} vs 2x{m}",
          "robustness")


# --------------------------------------------------------------- features ----
def validate_features(supp):
    means = supp["class_means"]
    if os.path.exists(os.path.join(RESULTS, "feature_means.csv")):
        rows = {r["feature"]: r for r in csv.DictReader(
            open(os.path.join(RESULTS, "feature_means.csv")))}
        ok = all(approx(means[f]["human_mean"], rows[f]["real_world_mean"], RAW_TOL)
                 and approx(means[f]["llm_mean"], rows[f]["ai_mean"], RAW_TOL)
                 for f in means if f in rows)
        check("feature-means CSV agrees with analysis JSON", ok, "", "features")

    ag = supp["importance_agreement"]
    check("importance overlap count is 7", ag["features_in_both_top20"] == 7,
          f"got {ag['features_in_both_top20']}", "features")
    check("Spearman rho is 0.00 with p = 1.00",
          approx(ag["spearman_rho"], 0.0, RPT_TOL) and approx(ag["p_value"], 1.0, RPT_TOL),
          f"rho={ag['spearman_rho']:.4f} p={ag['p_value']:.4f}", "features")


# -------------------------------------------------------------- checksums ----
def validate_checksums(summary):
    from scripts.build_corpus import sha256_file, NAZARIO_FILES, RAW
    stored = summary["provenance"]["sha256"]
    ok, detail = True, []
    for name, digest in stored.items():
        path = os.path.join(RAW, name)
        if not os.path.exists(path):
            detail.append(f"{name}: raw file absent (cannot verify)")
            continue
        actual = sha256_file(path)
        if actual != digest:
            ok = False
            detail.append(f"{name}: MISMATCH")
    check("source checksums agree with the reported values", ok,
          "; ".join(detail) or "all present files match", "checksums")


# ------------------------------------------------------------------ docx ----
def _num(s):
    m = re.search(r"-?\d+\.\d+|-?\d+", s.replace(",", ""))
    return float(m.group()) if m else None


def validate_docx(path, results, supp, grouped):
    """Compare the numbers actually printed in the dissertation with the
    freshly generated artefact outputs."""
    if not os.path.exists(path):
        check("dissertation .docx present for table verification", False,
              f"missing {path}", "document")
        return
    import docx
    d = docx.Document(path)

    def rows_of(table):
        return [[c.text.strip() for c in r.cells] for r in table.rows]

    tables = [rows_of(t) for t in d.tables]

    # --- Table 6.1: per-model random-split metrics ---
    t61 = next((t for t in tables if t and t[0][0].lower() == "model"
                and any("Precision" in c for c in t[0])), None)
    if t61 is None:
        check("Table 6.1 located in document", False, "", "document")
    else:
        name_map = {"logistic regression": "logreg", "random forest": "random_forest",
                    "xgboost": "xgboost"}
        ok, detail = True, []
        for row in t61[1:]:
            key = name_map.get(row[0].strip().lower().replace("\n", " "))
            if not key:
                continue
            r = results[key]
            for col, field in ((1, "precision"), (2, "recall"), (3, "f1"),
                               (5, "roc_auc"), (6, "accuracy")):
                doc_v = _num(row[col])
                if doc_v is None or not approx(doc_v, r[field], RPT_TOL):
                    ok = False
                    detail.append(f"{key}.{field}: doc={row[col]} artefact={r[field]:.3f}")
        check("Table 6.1 values agree with generated results", ok,
              "; ".join(detail) or "all match", "document")

    # --- Table 6.4: robustness conditions ---
    t64 = next((t for t in tables if t and "Evaluation" in t[0][0]
                and any("Logistic" in c for c in t[0])), None)
    if t64 is None:
        check("Table 6.4 located in document", False, "", "document")
    else:
        gs, s22, abl = (grouped["grouped_split"],
                        grouped["sensitivity_2022_only"], supp["ablation"])
        expected = {
            "random split": {k: results[k]["f1"] for k in
                             ("logreg", "random_forest", "xgboost")},
            "ablation": {k: abl[k]["f1"] for k in
                         ("logreg", "random_forest", "xgboost")},
            "grouped": {k: gs[k]["f1"] for k in
                        ("logreg", "random_forest", "xgboost")},
            "2022": {k: s22[k]["f1"] for k in
                     ("logreg", "random_forest", "xgboost")},
        }
        ok, detail = True, []
        for row in t64[1:]:
            label = row[0].lower()
            key = ("random split" if "random" in label else
                   "ablation" if "ablation" in label or "structural" in label else
                   "grouped" if "group" in label or "cluster" in label else
                   "2022" if "2022" in label else None)
            if key is None:
                continue
            for col, model in ((3, "logreg"), (4, "random_forest"), (5, "xgboost")):
                if col >= len(row):
                    continue
                doc_v = _num(row[col])
                if doc_v is None or not approx(doc_v, expected[key][model], RPT_TOL):
                    ok = False
                    detail.append(f"{key}/{model}: doc={row[col].splitlines()[0]} "
                                  f"artefact={expected[key][model]:.3f}")
        check("Table 6.4 values agree with generated robustness outputs", ok,
              "; ".join(detail) or "all match", "document")

    # --- feature means quoted in Table 6.2 ---
    means = supp["class_means"]
    label_map = {"flesch reading ease": "flesch_reading_ease",
                 "flesch–kincaid grade": "flesch_kincaid_grade",
                 "flesch-kincaid grade": "flesch_kincaid_grade",
                 "avg. word length": "avg_word_length",
                 "urgency-word ratio": "urgency_word_ratio",
                 "digit ratio": "digit_ratio",
                 "word count": "word_count"}
    t62 = next((t for t in tables if t and "Feature" in t[0][0]
                and any("mean" in c.lower() for c in t[0])), None)
    if t62 is not None:
        ok, detail = True, []
        for row in t62[1:]:
            key = label_map.get(row[0].strip().lower())
            if not key:
                continue
            for col, side in ((1, "human_mean"), (2, "llm_mean")):
                doc_v = _num(row[col])
                # tolerance follows the precision the document actually prints:
                # an integer such as "107" tolerates half a unit, "60.0" tolerates 0.05
                dec = len(row[col].split(".")[1].strip()) if "." in row[col] else 0
                tol = MEAN_TOL if dec else 0.5
                if doc_v is None or not approx(doc_v, means[key][side], tol):
                    ok = False
                    detail.append(f"{key}.{side}: doc={row[col]} "
                                  f"artefact={means[key][side]:.3f}")
        check("feature means in the document agree with the analysis", ok,
              "; ".join(detail) or "all match", "document")

    # --- document integrity ---
    body = "\n".join(p.text for p in d.paragraphs)
    check("document retains a table of contents", "Table of Contents" in body,
          "", "document")
    check("document retains List of Figures and List of Tables",
          "List of Figures" in body and "List of Tables" in body, "", "document")
    check("no unresolved placeholder text remains",
          "[Student ID]" not in body and "Right-click" not in body,
          "", "document")
    imgs = [r for r in d.part.rels.values() if "image" in r.reltype]
    check("figures are embedded in the document", len(imgs) >= 4,
          f"{len(imgs)} images", "document")


def main():
    docx_path = sys.argv[1] if len(sys.argv) > 1 else DOCX_DEFAULT
    summary = json.load(open(os.path.join("data", "corpus_summary.json")))
    results = {r["model"]: r for r in json.load(open(os.path.join(RESULTS, "results.json")))}
    supp = json.load(open(os.path.join(RESULTS, "supplementary.json")))
    grouped = json.load(open(os.path.join(RESULTS, "grouped_split.json")))

    validate_corpus(summary)
    n_test = validate_split(results)
    validate_models(results, n_test)
    validate_robustness(supp, grouped)
    validate_features(supp)
    validate_checksums(summary)
    validate_docx(os.path.abspath(docx_path), results, supp, grouped)

    failed = [c for c in checks if not c["pass"]]
    width = max(len(c["check"]) for c in checks) + 2
    print("\nFINAL ALIGNMENT VALIDATION")
    print("=" * (width + 10))
    current = None
    for c in checks:
        if c["group"] != current:
            current = c["group"]
            print(f"\n[{current.upper()}]")
        status = "PASS" if c["pass"] else "FAIL"
        print(f"  {c['check']:<{width}} {status}")
        if c["detail"] and not c["pass"]:
            print(f"      -> {c['detail']}")
    print("\n" + "=" * (width + 10))
    print(f"{len(checks) - len(failed)}/{len(checks)} checks passed"
          f"{'' if not failed else f' — {len(failed)} FAILED'}")

    report = {
        "validation_date": _dt.datetime.now().isoformat(timespec="seconds"),
        "environment": {
            "python": platform.python_version(),
            "platform": platform.platform(),
        },
        "corpus_checks": [c for c in checks if c["group"] == "corpus"],
        "split_checks": [c for c in checks if c["group"] == "split"],
        "model_checks": [c for c in checks if c["group"] == "models"],
        "robustness_checks": [c for c in checks if c["group"] == "robustness"],
        "feature_checks": [c for c in checks if c["group"] == "features"],
        "checksum_checks": [c for c in checks if c["group"] == "checksums"],
        "document_checks": [c for c in checks if c["group"] == "document"],
        "total_checks": len(checks),
        "passed": len(checks) - len(failed),
        "failed": len(failed),
        "overall": "PASS" if not failed else "FAIL",
    }
    os.makedirs(RESULTS, exist_ok=True)
    with open(os.path.join(RESULTS, "final_alignment_validation.json"), "w") as fh:
        json.dump(report, fh, indent=2)
    print(f"wrote {RESULTS}/final_alignment_validation.json")
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()
